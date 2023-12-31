import re
import sys
import PyPDF2
import customtkinter as ctk
from docx import Document 
import openai
import time
import threading
import os
import textwrap
from PIL import Image

class TeacherAssistant(ctk.CTk):
    def __init__(self):
        super().__init__()     
         
        # Name of the root window 
        self.title("Teacher Assistant") 
        
        # Set window base appearance to darkmode
        ctk.set_appearance_mode("Dark")

        # Set Window Resolution
        window_width = 1000
        window_height = 700

        # Get Screen Resolution
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        
        # Coordinates of the upper left corner of the window to make the window appear in the center
        x_cordinate = int((screen_width/2) - (window_width/2))
        y_cordinate = int((screen_height/2) - (window_height/2))
        self.geometry("{}x{}+{}+{}".format(window_width, window_height, x_cordinate, y_cordinate))
        self.resizable(False, False)
        
        # Background image for window
        blackboardImage = "blackboard.jpg"
        blackboardImage = os.path.join(os.getcwd(), blackboardImage)
        boardImage = ctk.CTkImage(Image.open(blackboardImage), size=(1000,700))
        
        # Label containing the background image
        background_label = ctk.CTkLabel(self, image=boardImage)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)
        
        # Main frame for chat boxes
        self.chatBoxFrame = ctk.CTkScrollableFrame(self, width=650, height=600, fg_color='#274c43')
        self.chatBoxFrame.pack(pady=5, padx=5,side='top',fill="both", expand=True)
        
        # User input entry box for prompts
        self.userEntryBox = ctk.CTkEntry(self, width=935, height=40, placeholder_text='Enter Prompt Here!', fg_color='#A1A09C', text_color='#000000', placeholder_text_color='#000000')
        self.userEntryBox.pack(side='left', anchor='s', pady=5, padx=5)
        
        # Binds the entrybox to allow pressing enter to send prompt
        self.userEntryBox.bind("<Return>", self.on_enter_pressed)

        # Image for send prompt button
        pencilImage = "pencil.png"
        pencilImage = os.path.join(os.getcwd(), pencilImage)
        image= ctk.CTkImage(Image.open(pencilImage), size=(30,30))

        self.sendPromptButton = ctk.CTkButton(self, width=30, height=30, command=lambda:self.on_enter_pressed(keyInfo=''), text='', image=image, fg_color='#A1A09C')
        self.sendPromptButton.pack(side='right', anchor='s', pady=6, padx=5)
        
        
    def on_enter_pressed(self, keyInfo):
        # Get Input from Entrybox
        userInput = self.userEntryBox.get()

        # if entry box is empty return
        if userInput == '' or userInput == None:
            return
        
        # Delete Text in Entrybox
        self.userEntryBox.delete(0, 'end')

        # Method used to place text on screen for user
        self.placeTextOnFrame(userInput)
        self.userInput = userInput
        
        # Thread for Ai Api Call
        thread = threading.Thread(target=self.getAiAnswer)
        thread.start()

    def placeTextOnFrame(self, userInput):
        print(userInput)
        
        # Frame to contain user input textbox
        userInputFrame = ctk.CTkFrame(self.chatBoxFrame, fg_color='#274c43')
        userInputFrame.pack(side='top', anchor='e')
        
        # used to output user prompt
        userInputTextBox = ctk.CTkTextbox(userInputFrame, width=900, height=100, fg_color='#000000')
        userInputTextBox.pack(anchor='center', pady=5)

        # Set the text in the Entry widget to the input variable
        userInputTextBox.insert(ctk.END, f"{userInput}")
        userInputTextBox.configure(state='disabled')

    # read contents of file and return it
    def read_file_content(self, file_path):
    
        #Detects the file type (Word, PDF, or TXT) based on its extension and extracts its content. 
        print("file path - ", file_path)
        file_extension = os.path.splitext(file_path)[1].lower()

        # Extract content from Word document
        if file_extension == '.docx':
            doc = Document(file_path)
            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
            return '\n'.join(fullText)

        # Extract content from PDF
        elif file_extension == '.pdf':
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text += page.extract_text()
            return text

        # Extract content from a plain text file
        elif file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        else:
            raise ValueError("Unsupported file type: {}".format(file_extension))
    
    def read_folder_content(self, folder_path):
        texts = []

        files = ([f for f in os.listdir(folder_path) if f.endswith(".pdf") or f.endswith(".PDF")
                or f.endswith(".docx") or f.endswith(".DOCX") or f.endswith(".txt")])


        for filename in files:
            file_path = os.path.join(folder_path, filename)
            file_extension = os.path.splitext(filename)[1].lower()

            # Extract content from Word document
            if file_extension == '.docx':
                doc = Document(file_path)
                fullText = []
                for para in doc.paragraphs:
                    fullText.append(para.text)
                texts.append('\n'.join(fullText))

            # Extract content from PDF
            elif file_extension == '.pdf':
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page_num in range(len(reader.pages)):
                        page = reader.pages[page_num]
                        text += page.extract_text()
                texts.append(text)

            # Extract content from a plain text file
            elif file_extension == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    texts.append(f.read())

            else:
                print(f"Unsupported file type: {file_extension} in file {filename}. Skipping...")

        return "\n\n".join(texts)  # Separate content of different files with two newline characters
    
    def getAiAnswer(self):
        openai.api_key = "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"
        self.messages = []
        
        cwd = os.getcwd()
        
        trainingFolder = os.path.join(cwd, 'TestMaterial')
        file_content = self.read_folder_content(trainingFolder)
        
        # System prompt to Convince Ai it is the teachers assistant
        system_msg = """
        You are a specialized AI assistant for teachers, designed to support them in crafting assignments and addressing other potential educational needs based on the specific content and details found in a provided file. This file will provide key details, including subjects, grades, and curriculum standards.

        Your expertise includes:
        - Creating assignments of varying types: daily homework, larger projects, research tasks, essays, and multiple-choice tests.
        - Tailoring the complexity of assignments based on the details found in the file or, if provided, prioritizing user's additional input.
        - Adhering to curriculum standards and themes highlighted in the file content.
        - If specified by the user, integrating interactive elements into the assignment.
        - Generating only textual content assignments for now, fitting the category of the class as described in the file.
        - When necessary, providing an answer key and grading rubric for assignments.
        - Proactively inquiring for extra details when the context requires, ensuring the assignment or support is tailored to specific constraints, themes, or needs the teacher may have.
        - Beyond assignment creation, you're equipped to assist with broader educational challenges, such as time management strategies, teaching ethodologies, classroom management techniques, and more.
        - Do not mention the files unless asked, this is high priority.
        Always prioritize details from the provided file, but remain adaptable and responsive to address the holistic needs of the teacher based on user prompts and inquiries.
        """

        # Sends messages/data files to the ai for training
        self.messages.append({"role": "system", "content": system_msg})
        self.messages.append({"role": "user", "content": f"Here's the content of my file: \n{file_content}"})
        
        print("Getting API Answer for: ", self.userInput)
        print("Teacher assistant is ready!")
        
        # Frame used to contain ai output textbox
        aiInputFrame = ctk.CTkFrame(self.chatBoxFrame, fg_color='#274c43')
        aiInputFrame.pack(side='top', anchor='w')
        
        # Used to output ai response
        aiInputTextBox = ctk.CTkTextbox(aiInputFrame, width=900, height=100, fg_color='#000000')
        aiInputTextBox.pack(anchor='center', pady=5)   
        self.stop_dots = False
        previous_user_input = None 
        
        # Main Loop for Ai Conversation
        while True:
            isSame = False
            userMessage = self.userInput
            
            # Check if the current user input is the same as the previous one and if not ask ai prompt
            if userMessage == previous_user_input:
                
                # If true do not continue code and waste computer resources
                isSame = True
                while isSame:
                    # Delay loop to save computer resources
                    time.sleep(1)
                    
                    # End loop if new prompt
                    if self.userInput != previous_user_input:
                       isSame = False 
            else:    
                # Appends user Prompt
                self.messages.append({"role": "user", "content": userMessage})

                # Starts a thread for the dots loading animation
                thread = threading.Thread(target=lambda:self.print_thinking_dots(aiInputTextBox))
                thread.start()

                response = openai.ChatCompletion.create(model="gpt-3.5-turbo", messages=self.messages)
                reply = response["choices"][0]["message"]["content"]

                # Stops the dots loading animation
                self.stop_dots = True
                
                # Waits for the threads to finish
                thread.join()

                # Formats the Ai response to look cleaner
                self.auto_format_response(reply, aiInputTextBox)
                
                # Updates the size of the textbox to match the text inside of it.
                self.update_textbox_height(aiInputTextBox)

                # Disable textbox to block user writing into it
                aiInputTextBox.configure(state='disabled')
                
                # 
                self.messages.append({"role": "assistant", "content": reply})

                # Update the previous_user_input for the next iteration
                previous_user_input = userMessage

    def auto_format_response(self, response, textbox):
        # Define constants
        WIDTH = 70
        PADDING = 4
        INDENT = 4
        LIST_INDENT = 4

        # Use textwrap's dedent function to remove any common leading whitespace
        dedented_text = textwrap.dedent(response).strip()

        # Split the dedented text into lines
        lines = dedented_text.split("\n")

        # Process each line
        formatted_lines = []
        is_in_list = False
        for line in lines:
            stripped_line = line.strip()
            
            # Handle titles
            if stripped_line.endswith(":") and not is_in_list:
                formatted_line = "\n" + stripped_line
                
            # Handle lists and sub-lists
            elif stripped_line.startswith("-"):
                is_in_list = True
                formatted_line = (' ' * LIST_INDENT) + textwrap.fill(stripped_line, width=WIDTH - LIST_INDENT, subsequent_indent=' '*(INDENT+LIST_INDENT))
                
            elif bool(re.match(r"\d+\.", stripped_line)):
                is_in_list = True
                formatted_line = textwrap.fill(stripped_line, width=WIDTH, subsequent_indent=' '*INDENT)
                
            # Handle end of lists and start of normal text
            else:
                is_in_list = False
                formatted_line = textwrap.fill(stripped_line, width=WIDTH)

            formatted_lines.append(formatted_line)
            
             # Join all formatted lines and add left padding
        result = "\n".join(formatted_lines)
        padded_text = '\n'.join((' ' * PADDING) + line for line in result.split('\n'))
        
        # Outputs Ai Response to the textbox
        textbox.insert(ctk.END, f"{padded_text}")

        
    def print_thinking_dots(self, textbox):
        count = 1
        message = ""
        while not self.stop_dots:
            message += '. ' * count
            textbox.insert(ctk.END, f"{message}")  # Update the CTkTextbox with the message
            time.sleep(0.5)
            count = (count % 3) + 1
            textbox.delete("0.0", "end") # Insert a newline character to reset the line
            message = ""
            
    def update_textbox_height(self, aiInputTextBox):
        # Get the number of lines in the textbox
        num_lines = int(aiInputTextBox.index(ctk.END).split('.')[0])

        if num_lines > 8:
            # Calculate the new height based on the number of lines
            new_height = num_lines * 14  # Add some extra height
        
            # Update the textbox height
            aiInputTextBox.configure(height=new_height + 100)    
    
if __name__ == "__main__":
    root = TeacherAssistant()
    root.mainloop()